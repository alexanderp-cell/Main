#!/usr/bin/env python3
"""Build JET Technic teardown & logistics agreements (EN/RU) with Calibri throughout."""
import io
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT = 'Calibri'
DIR = Path(__file__).resolve().parent
LOGO = DIR / 'jet_technic_logo.png'
FOOTER = (
    'Egitim Mahallesi, Fahrettin Kerim Gokay Caddesi, Ortaklar Ishani, '
    'N: 71/57, Kadikoy, 34722, Istanbul, Turkey  |  sales@jetechnic.com  |  jetechnic.com'
)

LEGAL = 'Hava Araclari Lojistik Sanayi ve Ticaret Anonim'
TRADE = 'JET TECHNIC'
ADDR = 'Egitim Mahallesi, Fahrettin Kerim Gokay Caddesi, Ortaklar Is Merkezi, No: 71/57, Kadikoy, Istanbul, Turkey'
TEL = '+90 533 736 84 00'
DG = 'Serdar Guler'
CO_NO = '389609-5'
TAX = '484 212 4015'
BANK = 'YAPI KREDI BANK'
IBAN_USD = 'TR850006701000000015176740'
SWIFT = 'YAPITRISXXX'
EMAIL = 'sales@jetechnic.com'

# Prices intentionally left blank for fill-in (not taken from KP yet)
PRICE_ITEMS_EN = [
    'Aircraft Teardown per project scope — USD [●]',
    'Parts Shipment (incl. powerplant, LDGs, Cabin) — USD [●]',
]
PRICE_ITEMS_RU = [
    'Разбор воздушного судна (по объёму проекта) — USD [●]',
    'Отправка / логистика комплектующих (включая силовую установку, шасси, кабину) — USD [●]',
]

PARTS = [
    ('7121-19971-01AC', 'Cabin Pressure controller', '2'),
    ('5145-1-64', 'Audio Selector Panel', '1'),
    ('172625-7', 'Valve Assy PNEU', '2'),
    ('2233000-23-A', 'Digital Flight data A.C. Unit', '1'),
    ('69-78214-4', 'Aural Warning Module ASSEMBLY', '1'),
    ('2606672-4', 'Brake - Main Wheel', '1'),
    ('171497-05-01', 'Flight Management computer', '1'),
    ('822-1338-205', 'ATC Transponder', '2'),
    ('HG1050AE09', 'Unit Assy - Inertial Reference', '1'),
    ('HG1050AE10', 'Unit Assy - Inertial Reference', '1'),
    ('3605812-22', 'Starter Assy - ELEC', '1'),
    ('3617750-1', 'Control Assy Fuel', '1'),
    ('5-89354-3150', 'Windshield N1 Assy Co-Pilot', '1'),
    ('314A1502-58', 'Sleeve Assy - Turb Exch Primary', '1'),
    ('129666-2', 'Sensor-Precooler Control Valve', '2'),
    ('3289562-5', 'Valve Fan Air Precooler CTRL', '2'),
]

HL_RE = re.compile(
    r'(_{4,}|«____»|\[●\]|\[CLIENT FULL LEGAL NAME\]|\[ПОЛНОЕ НАИМЕНОВАНИЕ ЗАКАЗЧИКА\]|'
    r'\[full name, title\]|\[ФИО, должность\]|\[Charter / Power of Attorney[^\]]*\]|'
    r'\[Устава / доверенности[^\]]*\]|\[Name\]|\[ФИО\]|'
    r'\[delivery location\(s\)\]|\[место\(а\) поставки\])'
)

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
A = '{http://schemas.openxmlformats.org/drawingml/2006/main}'


def set_run_font(run, *, size=10, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    r = run._element.get_or_add_rPr()
    rFonts = r.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        r.insert(0, rFonts)
    for key in ('asciiTheme', 'hAnsiTheme', 'cstheme', 'eastAsiaTheme'):
        attr = qn(f'w:{key}')
        if attr in rFonts.attrib:
            del rFonts.attrib[attr]
    for key, val in (
        (qn('w:ascii'), FONT),
        (qn('w:hAnsi'), FONT),
        (qn('w:cs'), FONT),
        (qn('w:eastAsia'), FONT),
    ):
        rFonts.set(key, val)


def add_text(p, text, *, bold=False, size=10, yellow=False):
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if yellow:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return run


def write_runs(p, text, *, bold=False, size=10, yellow_all=False):
    if yellow_all or HL_RE.search(text):
        pos = 0
        for m in HL_RE.finditer(text):
            if m.start() > pos:
                add_text(p, text[pos:m.start()], bold=bold, size=size)
            add_text(p, m.group(), bold=bold, size=size, yellow=True)
            pos = m.end()
        if pos < len(text):
            add_text(p, text[pos:], bold=bold, size=size)
    else:
        add_text(p, text, bold=bold, size=size)


def add_para(doc, text, **kw):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(kw.get('sb', 0))
    pf.space_after = Pt(kw.get('sa', 2))
    pf.line_spacing = 1.05
    align = kw.get('align', 'justify')
    p.alignment = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
    }.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    write_runs(
        p, text,
        bold=kw.get('bold', False),
        size=kw.get('size', 10),
        yellow_all=kw.get('yellow_all', False),
    )
    return p


def add_mixed(doc, parts, **kw):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(kw.get('sa', 2))
    p.paragraph_format.line_spacing = 1.05
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold in parts:
        write_runs(p, text, bold=bold)
    return p


def setup_doc(doc):
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(10)
    for s in doc.sections:
        s.top_margin = Cm(1.2)
        s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(1.8)
        s.right_margin = Cm(1.5)
        s.header_distance = Cm(0.4)
        s.footer_distance = Cm(0.4)
    sec = doc.sections[0]
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.clear()
    hp.add_run().add_picture(str(LOGO), width=Cm(10))
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run(FOOTER)
    set_run_font(r, size=7.5)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def fill_cell(cell, title, lines):
    cell.text = ''
    p = cell.paragraphs[0]
    add_text(p, title, bold=True, size=9.5)
    for line in lines:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.02
        write_runs(p2, line, size=8.5)


def set_table_fonts(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=run.font.size.pt if run.font.size else 8.5, bold=bool(run.bold))


def add_parts_table(doc, lang='en'):
    table = doc.add_table(rows=1 + len(PARTS), cols=3)
    table.style = 'Table Grid'
    headers = (
        ('Part Number', 'Description', 'QTY') if lang == 'en'
        else ('Part Number', 'Наименование', 'Кол-во')
    )
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        add_text(p, h, bold=True, size=8.5)
    for ri, (pn, desc, qty) in enumerate(PARTS, start=1):
        for ci, val in enumerate((pn, desc, qty)):
            cell = table.rows[ri].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            add_text(p, val, size=8)
    set_table_fonts(table)
    return table


def patch_rfonts_element(rfonts, font=FONT):
    theme_keys = [k for k in list(rfonts.attrib.keys()) if 'Theme' in k or k.endswith('theme')]
    for k in theme_keys:
        del rfonts.attrib[k]
    for tag in ('ascii', 'hAnsi', 'cs', 'eastAsia'):
        rfonts.set(f'{W}{tag}', font)


def patch_xml_root(root, font=FONT):
    for rfonts in root.iter(f'{W}rFonts'):
        patch_rfonts_element(rfonts, font)
    for r in root.iter(f'{W}r'):
        texts = ''.join(t.text or '' for t in r.iter(f'{W}t'))
        if not texts:
            continue
        rpr = r.find(f'{W}rPr')
        if rpr is None:
            rpr = ET.Element(f'{W}rPr')
            r.insert(0, rpr)
        rf = rpr.find(f'{W}rFonts')
        if rf is None:
            rf = ET.SubElement(rpr, f'{W}rFonts')
        patch_rfonts_element(rf, font)
    for latin in root.iter(f'{A}latin'):
        latin.set('typeface', font)


def patch_docx_fonts(path: Path, font: str = FONT):
    replacements = (
        ('typeface="Cambria"', f'typeface="{font}"'),
        ('typeface="Times New Roman"', f'typeface="{font}"'),
        ('w:ascii="Times New Roman"', f'w:ascii="{font}"'),
        ('w:hAnsi="Times New Roman"', f'w:hAnsi="{font}"'),
        ('w:cs="Times New Roman"', f'w:cs="{font}"'),
        ('w:ascii="Cambria"', f'w:ascii="{font}"'),
        ('w:hAnsi="Cambria"', f'w:hAnsi="{font}"'),
        ('w:cs="Cambria"', f'w:cs="{font}"'),
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(path, 'r') as zin, zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith('word/') and item.filename.endswith('.xml'):
                text = data.decode('utf-8')
                for old, new in replacements:
                    text = text.replace(old, new)
                root = ET.fromstring(text.encode('utf-8'))
                patch_xml_root(root, font)
                data = ET.tostring(root, encoding='utf-8', xml_declaration=True)
            zout.writestr(item, data)
    path.write_bytes(buf.getvalue())


def build_en():
    doc = Document()
    setup_doc(doc)

    add_para(doc, 'SERVICE AGREEMENT No. ____ / JT-TD-2026', bold=True, align='center', size=12, sa=1)
    add_para(doc, 'Aircraft Teardown and Logistics Services', bold=True, align='center', size=10.5, sa=1)
    add_para(doc, 'Boeing 737-500, Registration VQ-BOO', bold=True, align='center', size=10, sa=3)
    add_para(doc, 'Place of conclusion: ________________    Date: «____» ____________ 2026', align='center', size=9.5, sa=5)

    add_mixed(doc, [
        (f'{TRADE} / {LEGAL}', True),
        (f' ({ADDR}; Company No. {CO_NO}; Tax No. {TAX}; tel. {TEL}; {EMAIL}; the "', False),
        ('Contractor', True),
        (f'"), represented by {DG}, Director General, on the one part, and', False),
    ], sa=3)
    add_mixed(doc, [
        ('[CLIENT FULL LEGAL NAME]', True),
        (' ([●]; reg. / tax ID: [●]; the "', False),
        ('Customer', True),
        ('"), represented by [full name, title] under [Charter / Power of Attorney], on the other part,', False),
    ], sa=3)
    add_para(doc, 'have agreed as follows:', sa=4)

    add_para(doc, '1. SUBJECT', bold=True, size=10.5, sa=2, align='left')
    add_para(doc, '1.1. The Contractor shall perform, and the Customer shall accept and pay for, aircraft teardown and logistics services for Boeing 737-500 VQ-BOO (the "Aircraft"), as follows:')
    add_para(doc, '(a) Preparation (fluid draining, power management, access and security, etc.);')
    add_para(doc, '(b) Teardown / parting-out: powerplant (2× CFM56-3 and APU), cabin, landing gears, and parts requested by the Customer as listed in Appendix No. 1;')
    add_para(doc, '(c) Packaging and crating (including engine stands);')
    add_para(doc, '(d) Aircraft utilization; and')
    add_para(doc, '(e) Shipment of material to [delivery location(s)].')
    add_para(doc, '1.2. Location of performance: JAT Tehnika, Belgrade, Serbia. The Customer shall provide airport and Aircraft access.')
    add_para(doc, '1.3. Parts storage at the teardown location is included in the Services. The final workscope for teardown shall be confirmed after Aircraft inspection; Appendix No. 1 may be updated by written agreement of the Parties.')
    add_para(doc, '1.4. Appendix No. 1 (Parts List) forms an integral part of this Agreement.')

    add_para(doc, '2. TERM', bold=True, size=10.5, sb=2, sa=2, align='left')
    add_para(doc, '2.1. Approximate duration of teardown: 60 (sixty) days from the Effective Date and provision of access and required documentation / authorizations by the Customer. Shipment timing depends on the final parts list and delivery locations.')
    add_para(doc, '2.2. The Customer shall accept completed stages / deliverables within 5 (five) business days or submit reasoned objections; otherwise they are deemed accepted.')

    add_para(doc, '3. PRICE AND PAYMENT', bold=True, size=10.5, sb=2, sa=2, align='left')
    add_para(doc, '3.1. The total Price (exclusive of taxes and fees) shall be agreed in writing and filled in below:')
    for item in PRICE_ITEMS_EN:
        add_para(doc, f'• {item}', align='left', sa=1)
    add_para(doc, 'Total: USD [●].', bold=True, sa=2, align='left')
    add_para(doc, '3.2. The teardown price is subject to confirmation of the final workscope after Aircraft inspection. Shipment costs are approximate and may be adjusted according to the final material list and delivery locations.')
    add_para(doc, '3.3. Payment option (select one; strike out the others):', sa=2)
    add_para(doc, '(A) 100% post-payment within 5 banking days after acceptance;', yellow_all=True, sa=1)
    add_para(doc, '(B) 50% on invoice and 50% within 5 banking days after acceptance; or', yellow_all=True, sa=1)
    add_para(doc, '(C) 100% advance within 3 banking days of invoice.', yellow_all=True, sa=1)
    add_para(doc, 'If not selected otherwise in writing, option (A) applies.', yellow_all=True, sa=2)
    add_para(doc, '3.4. Payment in USD to the Contractor\'s account in Clause 7. Remitting bank charges are for the Customer.')
    add_para(doc, '3.5. Extra work outside the agreed scope requires prior written agreement and separate pricing.')

    add_para(doc, '4. OBLIGATIONS', bold=True, size=10.5, sb=2, sa=2, align='left')
    add_para(doc, '4.1. The Contractor shall perform the Services professionally, comply with site safety rules and JAT Tehnika procedures, and deliver packaging/crating and shipment documentation as applicable.')
    add_para(doc, '4.2. The Customer shall provide timely access, coordination with JAT Tehnika, authorizations, technical documentation, delivery instructions, payment, and a contact person.')
    add_para(doc, '4.3. Title to removed parts and risk of loss shall pass as agreed in writing for each shipment. Until shipment, parts remain stored at the teardown location at the Contractor\'s care as included in the Services.')

    add_para(doc, '5. LIABILITY AND FORCE MAJEURE', bold=True, size=10.5, sb=2, sa=2, align='left')
    add_para(doc, '5.1. Except for willful misconduct or gross negligence, the Contractor\'s liability is limited to the Price actually paid. Neither Party is liable for indirect or consequential losses.')
    add_para(doc, '5.2. Force majeure excuses performance while it continues, subject to prompt notice. If it exceeds 30 days, either Party may terminate without penalty, paying for Services already performed.')

    add_para(doc, '6. GENERAL', bold=True, size=10.5, sb=2, sa=2, align='left')
    add_para(doc, '6.1. This Agreement enters into force on signing (the "Effective Date") and remains valid until fully performed.')
    add_para(doc, '6.2. Either Party may terminate for material breach not cured within 10 business days. The Customer may cancel before commencement subject to reimbursement of documented costs.')
    add_para(doc, '6.3. Amendments require written form signed by both Parties (including PDF scans of signed originals), including updates to Appendix No. 1 and the Price.')
    add_para(doc, '6.4. Confidential information shall be protected for 3 years after termination.')
    add_para(doc, '6.5. Notices by email to Clause 7 addresses are deemed received on the next business day.')
    add_para(doc, '6.6. Governed by the laws of Türkiye. Disputes not settled within 15 business days shall be submitted to the courts of Istanbul, Türkiye (English unless otherwise agreed).')

    add_para(doc, '7. DETAILS AND SIGNATURES', bold=True, size=10.5, sb=2, sa=3, align='left')
    t = doc.add_table(rows=2, cols=2)
    fill_cell(t.rows[0].cells[0], f'CONTRACTOR: {TRADE}', [
        LEGAL, ADDR, f'Company No.: {CO_NO}   Tax No.: {TAX}',
        f'Bank: {BANK}', f'IBAN (USD): {IBAN_USD}', f'SWIFT: {SWIFT}',
        f'Tel.: {TEL}   Email: {EMAIL}',
    ])
    fill_cell(t.rows[0].cells[1], 'CUSTOMER: [●]', [
        'Name: [●]', 'Address: [●]', 'Reg. / Tax ID: [●]',
        'Bank / IBAN / SWIFT: [●]', 'Email: [●]   Tel.: [●]',
    ])
    fill_cell(t.rows[1].cells[0], '', ['', f'________________ / {DG} /', 'Director General', 'L.S.'])
    fill_cell(t.rows[1].cells[1], '', ['', '________________ / [Name] /', 'Title: [●]', 'L.S.'])

    doc.add_page_break()
    add_para(doc, 'APPENDIX No. 1', bold=True, align='center', size=12, sa=2)
    add_para(doc, 'to Service Agreement No. ____ / JT-TD-2026', bold=True, align='center', size=10, sa=2)
    add_para(doc, 'PARTS LIST (as requested by the Customer)', bold=True, align='center', size=11, sa=4)
    add_para(doc, 'Aircraft: Boeing 737-500, registration mark VQ-BOO', align='center', size=10, sa=1)
    add_para(doc, 'Place of performance: JAT Tehnika, Belgrade, Republic of Serbia', align='center', size=10, sa=4)
    add_para(doc, '1. Parts to be removed / prepared for shipment:', bold=True, sa=3, align='left')
    add_parts_table(doc, 'en')
    add_para(doc, '', sa=3)
    add_para(doc, '2. Notes:', bold=True, sa=2, align='left')
    add_para(doc, '2.1. This list may be amended by written agreement of the Parties.')
    add_para(doc, '2.2. Final workscope and Price for teardown / shipment shall be confirmed after Aircraft inspection and based on the final material list and delivery locations.')
    add_para(doc, '2.3. This Appendix No. 1 is an integral part of the Agreement.', sa=6)
    add_para(doc, 'CONTRACTOR: ________________ / Serdar Guler /', align='left', sa=2)
    add_para(doc, 'CUSTOMER:   ________________ / [Name] /', align='left', sa=4)
    add_para(doc, 'DRAFT — not legal advice. Prices to be filled in.', size=8, sa=0, align='left')
    return doc


def build_ru():
    doc = Document()
    setup_doc(doc)

    add_para(doc, 'ДОГОВОР ВОЗМЕЗДНОГО ОКАЗАНИЯ УСЛУГ № ____ / JT-TD-2026', bold=True, align='center', size=12, sa=1)
    add_para(doc, 'разбор воздушного судна и логистические услуги', bold=True, align='center', size=10.5, sa=1)
    add_para(doc, 'Boeing 737-500, регистрационный знак VQ-BOO', bold=True, align='center', size=10, sa=3)
    add_para(doc, 'Место заключения: ________________    Дата: «____» ____________ 2026 г.', align='center', size=9.5, sa=5)

    add_mixed(doc, [
        (f'{TRADE} / {LEGAL}', True),
        (f' ({ADDR}; Company No. {CO_NO}; Tax No. {TAX}; тел. {TEL}; {EMAIL}; «', False),
        ('Исполнитель', True),
        (f'»), в лице Генерального директора {DG}, с одной стороны, и', False),
    ], sa=3)
    add_mixed(doc, [
        ('[ПОЛНОЕ НАИМЕНОВАНИЕ ЗАКАЗЧИКА]', True),
        (' ([●]; рег. / налог. номер: [●]; «', False),
        ('Заказчик', True),
        ('»), в лице [ФИО, должность] на основании [Устава / доверенности], с другой стороны,', False),
    ], sa=3)
    add_para(doc, 'договорились о нижеследующем:', sa=4)

    add_para(doc, '1. ПРЕДМЕТ', bold=True, size=10.5, sa=2, align='left')
    add_para(doc, '1.1. Исполнитель оказывает, а Заказчик принимает и оплачивает услуги по разбору воздушного судна и логистике в отношении Boeing 737-500 VQ-BOO (далее — «ВС»), а именно:')
    add_para(doc, '(a) подготовка (слив жидкостей, управление питанием, доступ и безопасность и т.п.);')
    add_para(doc, '(b) разбор / parting-out: силовая установка (2× CFM56-3 и ВСУ), кабина, шасси, а также комплектующие по запросу Заказчика согласно Приложению № 1;')
    add_para(doc, '(c) упаковка и крейтинг (включая стенды для двигателей);')
    add_para(doc, '(d) утилизация остатков ВС; и')
    add_para(doc, '(e) отправка материалов в [место(а) поставки].')
    add_para(doc, '1.2. Место оказания: JAT Tehnika, Белград, Сербия. Заказчик обеспечивает доступ в аэропорт и к ВС.')
    add_para(doc, '1.3. Хранение снятых комплектующих на площадке разбора включено в Услуги. Окончательный объём разбора подтверждается после инспекции ВС; Приложение № 1 может быть изменено по письменному соглашению Сторон.')
    add_para(doc, '1.4. Приложение № 1 (Перечень комплектующих) является неотъемлемой частью Договора.')

    add_para(doc, '2. СРОК', bold=True, size=10.5, sb=2, sa=2, align='left')
    add_para(doc, '2.1. Ориентировочный срок разбора: 60 (шестьдесят) дней с даты вступления Договора в силу и предоставления Заказчиком доступа и необходимой документации / разрешений. Сроки отправки зависят от финального перечня и мест поставки.')
    add_para(doc, '2.2. Заказчик принимает завершённые этапы / результаты в течение 5 (пяти) рабочих дней либо направляет мотивированные возражения; иначе они считаются принятыми.')

    add_para(doc, '3. ЦЕНА И ОПЛАТА', bold=True, size=10.5, sb=2, sa=2, align='left')
    add_para(doc, '3.1. Общая Цена (без учёта налогов и сборов) согласовывается письменно и указывается ниже:')
    for item in PRICE_ITEMS_RU:
        add_para(doc, f'• {item}', align='left', sa=1)
    add_para(doc, 'Итого: USD [●].', bold=True, sa=2, align='left')
    add_para(doc, '3.2. Цена разбора подлежит подтверждению после инспекции ВС и уточнения окончательного объёма работ. Стоимость логистики ориентировочная и может быть скорректирована по финальному перечню и местам поставки.')
    add_para(doc, '3.3. Вариант оплаты (выбрать один; остальные вычеркнуть):', sa=2)
    add_para(doc, '(A) 100% постоплата в течение 5 банковских дней после приёмки;', yellow_all=True, sa=1)
    add_para(doc, '(B) 50% по счёту и 50% в течение 5 банковских дней после приёмки; либо', yellow_all=True, sa=1)
    add_para(doc, '(C) 100% предоплата в течение 3 банковских дней с даты счёта.', yellow_all=True, sa=1)
    add_para(doc, 'Если иное не выбрано письменно, применяется вариант (A).', yellow_all=True, sa=2)
    add_para(doc, '3.4. Оплата в USD на счёт Исполнителя (раздел 7). Комиссии банка-отправителя — за счёт Заказчика.')
    add_para(doc, '3.5. Работы вне согласованного объёма — только по предварительному письменному согласованию и отдельной цене.')

    add_para(doc, '4. ОБЯЗАННОСТИ', bold=True, size=10.5, sb=2, sa=2, align='left')
    add_para(doc, '4.1. Исполнитель оказывает Услуги профессионально, соблюдает правила ТБ и процедуры JAT Tehnika, обеспечивает упаковку/крейтинг и сопроводительную документацию по отправке (при применимости).')
    add_para(doc, '4.2. Заказчик своевременно обеспечивает доступ, координацию с JAT Tehnika, разрешения, техническую документацию, инструкции по поставке, оплату и контактное лицо.')
    add_para(doc, '4.3. Переход права собственности и рисков по снятым комплектующим определяется письменно для каждой отправки. До отправки комплектующие хранятся на площадке разбора в рамках Услуг.')

    add_para(doc, '5. ОТВЕТСТВЕННОСТЬ И ФОРС-МАЖОР', bold=True, size=10.5, sb=2, sa=2, align='left')
    add_para(doc, '5.1. За исключением умысла или грубой неосторожности ответственность Исполнителя ограничена фактически уплаченной Ценой. Стороны не отвечают за косвенные убытки.')
    add_para(doc, '5.2. Форс-мажор освобождает от ответственности на период его действия при своевременном уведомлении. При продолжении более 30 дней любая Сторона вправе расторгнуть Договор без штрафов с оплатой фактически оказанных Услуг.')

    add_para(doc, '6. ОБЩИЕ ПОЛОЖЕНИЯ', bold=True, size=10.5, sb=2, sa=2, align='left')
    add_para(doc, '6.1. Договор вступает в силу с даты подписания и действует до полного исполнения.')
    add_para(doc, '6.2. Любая Сторона вправе расторгнуть Договор при существенном нарушении, не устранённом в течение 10 рабочих дней. Заказчик вправе отказаться до начала работ с возмещением документально подтверждённых расходов Исполнителя.')
    add_para(doc, '6.3. Изменения — только в письменной форме, подписанной обеими Сторонами (включая PDF-сканы), в том числе обновления Приложения № 1 и Цены.')
    add_para(doc, '6.4. Конфиденциальная информация защищается в течение 3 лет после прекращения Договора.')
    add_para(doc, '6.5. Уведомления по email (раздел 7) считаются полученными на следующий рабочий день.')
    add_para(doc, '6.6. Применимое право — Турция. Споры, не урегулированные за 15 рабочих дней, передаются в суды Стамбула (язык — английский, если не согласовано иное).')

    add_para(doc, '7. РЕКВИЗИТЫ И ПОДПИСИ', bold=True, size=10.5, sb=2, sa=3, align='left')
    t = doc.add_table(rows=2, cols=2)
    fill_cell(t.rows[0].cells[0], f'ИСПОЛНИТЕЛЬ: {TRADE}', [
        LEGAL, ADDR, f'Company No.: {CO_NO}   Tax No.: {TAX}',
        f'Банк: {BANK}', f'IBAN (USD): {IBAN_USD}', f'SWIFT: {SWIFT}',
        f'Тел.: {TEL}   Email: {EMAIL}',
    ])
    fill_cell(t.rows[0].cells[1], 'ЗАКАЗЧИК: [●]', [
        'Наименование: [●]', 'Адрес: [●]', 'Рег. / налог. номер: [●]',
        'Банк / IBAN / SWIFT: [●]', 'Email: [●]   Тел.: [●]',
    ])
    fill_cell(t.rows[1].cells[0], '', ['', f'________________ / {DG} /', 'Генеральный директор', 'М.П.'])
    fill_cell(t.rows[1].cells[1], '', ['', '________________ / [ФИО] /', 'Должность: [●]', 'М.П.'])

    doc.add_page_break()
    add_para(doc, 'ПРИЛОЖЕНИЕ № 1', bold=True, align='center', size=12, sa=2)
    add_para(doc, 'к Договору возмездного оказания услуг № ____ / JT-TD-2026', bold=True, align='center', size=10, sa=2)
    add_para(doc, 'ПЕРЕЧЕНЬ КОМПЛЕКТУЮЩИХ (по запросу Заказчика)', bold=True, align='center', size=11, sa=4)
    add_para(doc, 'Воздушное судно: Boeing 737-500, регистрационный знак VQ-BOO', align='center', size=10, sa=1)
    add_para(doc, 'Место оказания: JAT Tehnika, Белград, Республика Сербия', align='center', size=10, sa=4)
    add_para(doc, '1. Комплектующие к снятию / подготовке к отправке:', bold=True, sa=3, align='left')
    add_parts_table(doc, 'ru')
    add_para(doc, '', sa=3)
    add_para(doc, '2. Примечания:', bold=True, sa=2, align='left')
    add_para(doc, '2.1. Перечень может быть изменён по письменному соглашению Сторон.')
    add_para(doc, '2.2. Окончательный объём и Цена разбора / логистики подтверждаются после инспекции ВС с учётом финального перечня и мест поставки.')
    add_para(doc, '2.3. Настоящее Приложение № 1 является неотъемлемой частью Договора.', sa=6)
    add_para(doc, 'ИСПОЛНИТЕЛЬ: ________________ / Serdar Guler /', align='left', sa=2)
    add_para(doc, 'ЗАКАЗЧИК:    ________________ / [ФИО] /', align='left', sa=4)
    add_para(doc, 'ЧЕРНОВИК — не юридическая консультация. Цены подлежат заполнению.', size=8, sa=0, align='left')
    return doc


def main():
    outputs = [
        (build_en(), DIR / 'Service_Agreement_JT-TD-2026_EN.docx'),
        (build_ru(), DIR / 'Dogovor_JT-TD-2026_RU.docx'),
    ]
    for doc, path in outputs:
        doc.save(path)
        patch_docx_fonts(path)
        print('saved', path.name)


if __name__ == '__main__':
    main()
