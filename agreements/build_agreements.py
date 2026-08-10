#!/usr/bin/env python3
"""Build JET Technic inspection agreements (EN/RU) with Calibri throughout."""
import io
import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT = 'Calibri'
DIR = Path(__file__).resolve().parent
LOGO = DIR / 'jet_technic_logo.png'
FOOTER = (
    'Egitim Mahallesi, Fahrettin Kerim Gokay Caddesi, Ortaklar Ishani, '
    'N: 71/57, Kadikoy, 34722, Istanbul, Turkey  |  sales@jetechnic.com  |  jetechnic.com'
)

LEGAL = 'Hava Araclari Lojistik Sanayi ve Ticaret Anonim'
TRADE = 'JET TECHNIC'
ADDR = 'Egitim Mahallesi, Fahrettin Kerim Gokay Caddesi, Ortaklar Is Merkezi, No: 71/57, Kadikoy, Istanbul, Turkey'
TEL = '+90 533 736 84 00'
DG = 'Serdar Guler'
CO_NO = '389609-5'
TAX = '484 212 4015'
BANK = 'YAPI KREDI BANK'
IBAN_USD = 'TR850006701000000015176740'
SWIFT = 'YAPITRISXXX'
EMAIL = 'sales@jetechnic.com'

ITEMS_EN = [
    ('BSI CFM56-3 – 2 Engines (scope TBC)', 'USD 9,500.00'),
    ('BSI APU APS2000', 'USD 4,000.00'),
    ('Aircraft GVI (landing gears, cabin, avionics)', 'USD 3,000.00'),
    ('Paperwork inspection (Engines, APU)', 'USD 3,500.00'),
    ('Transportation and accommodation', 'USD 4,000.00'),
]
ITEMS_RU = [
    ('BSI CFM56-3 — 2 двигателя (объём уточняется)', 'USD 9 500,00'),
    ('BSI ВСУ APS2000', 'USD 4 000,00'),
    ('GVI ВС (шасси, кабина, авионика)', 'USD 3 000,00'),
    ('Проверка документации (двигатели, ВСУ)', 'USD 3 500,00'),
    ('Транспорт и проживание', 'USD 4 000,00'),
]

HL_RE = re.compile(
    r'(_{4,}|«____»|\[●\]|\[CLIENT FULL LEGAL NAME\]|\[ПОЛНОЕ НАИМЕНОВАНИЕ ЗАКАЗЧИКА\]|'
    r'\[full name, title\]|\[ФИО, должность\]|\[Charter / Power of Attorney[^\]]*\]|'
    r'\[Устава / доверенности[^\]]*\]|\[Name\]|\[ФИО\])'
)

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
A = '{http://schemas.openxmlformats.org/drawingml/2006/main}'


def set_run_font(run, *, size=10, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    r = run._element.get_or_add_rPr()
    rFonts = r.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = run._element.makeelement(qn('w:rFonts'), {})
        r.insert(0, rFonts)
    for key in ('asciiTheme', 'hAnsiTheme', 'cstheme', 'eastAsiaTheme'):
        attr = qn(f'w:{key}')
        if attr in rFonts.attrib:
            del rFonts.attrib[attr]
    for key, val in (
        (qn('w:ascii'), FONT),
        (qn('w:hAnsi'), FONT),
        (qn('w:cs'), FONT),
        (qn('w:eastAsia'), FONT),
    ):
        rFonts.set(key, val)


def add_text(p, text, *, bold=False, size=10, yellow=False):
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if yellow:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return run


def write_runs(p, text, *, bold=False, size=10, yellow_all=False):
    if yellow_all or HL_RE.search(text):
        pos = 0
        for m in HL_RE.finditer(text):
            if m.start() > pos:
                add_text(p, text[pos:m.start()], bold=bold, size=size)
            add_text(p, m.group(), bold=bold, size=size, yellow=True)
            pos = m.end()
        if pos < len(text):
            add_text(p, text[pos:], bold=bold, size=size)
    else:
        add_text(p, text, bold=bold, size=size)


def add_para(doc, text, **kw):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(kw.get('sb', 0))
    pf.space_after = Pt(kw.get('sa', 2))
    pf.line_spacing = 1.05
    align = kw.get('align', 'justify')
    p.alignment = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'left': WD_ALIGN_PARAGRAPH.LEFT,
    }.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
    write_runs(p, text, bold=kw.get('bold', False), size=kw.get('size', 10), yellow_all=kw.get('yellow_all', False))
    return p


def add_mixed(doc, parts, **kw):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(kw.get('sa', 2))
    p.paragraph_format.line_spacing = 1.05
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold in parts:
        write_runs(p, text, bold=bold)
    return p


def setup_doc(doc):
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(10)
    for s in doc.sections:
        s.top_margin = Cm(1.2)
        s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(1.8)
        s.right_margin = Cm(1.5)
        s.header_distance = Cm(0.4)
        s.footer_distance = Cm(0.4)
    sec = doc.sections[0]
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.clear()
    hp.add_run().add_picture(str(LOGO), width=Cm(10))
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run(FOOTER)
    set_run_font(r, size=7.5)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def fill_cell(cell, title, lines):
    cell.text = ''
    p = cell.paragraphs[0]
    add_text(p, title, bold=True, size=9.5)
    for line in lines:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        p2.paragraph_format.line_spacing = 1.02
        write_runs(p2, line, size=8.5)


def build_en():
    doc = Document()
    setup_doc(doc)
    add_para(doc, 'SERVICE AGREEMENT No. ____ / JT-BSI-2026', bold=True, align='center', size=12, sa=1)
    add_para(doc, 'Aircraft Inspection and Engine/APU Borescope Inspection', bold=True, align='center', size=10.5, sa=1)
    add_para(doc, 'Boeing 737-500, Registration VQ-BOO', bold=True, align='center', size=10, sa=3)
    add_para(doc, 'Place of conclusion: ________________    Date: «____» ____________ 2026', align='center', size=9.5, sa=5)
    add_mixed(doc, [
        (f'{TRADE} / {LEGAL}', True),
        (f' ({ADDR}; Company No. {CO_NO}; Tax No. {TAX}; tel. {TEL}; {EMAIL}; the "', False),
        ('Contractor', True),
        (f'"), represented by {DG}, Director General, on the one part, and', False),
    ], sa=3)
    add_mixed(doc, [
        ('[CLIENT FULL LEGAL NAME]', True),
        (' ([●]; reg. / tax ID: [●]; the "', False),
        ('Customer', True),
        ('"), represented by [full name, title] under [Charter / Power of Attorney], on the other part,', False),
    ], sa=3)
    add_para(doc, 'have agreed as follows:', sa=4)

    add_para(doc, '1. SUBJECT', bold=True, size=10.5, sa=2, align='left')
    add_para(doc, '1.1. The Contractor shall perform, and the Customer shall accept and pay for, the following services for Boeing 737-500 VQ-BOO (the "Aircraft"):')
    add_para(doc, '(a) BSI of two (2) CFM56-3 engines (scope to be confirmed); (b) BSI of APU APS2000; (c) Aircraft GVI (landing gears, cabin, avionics compartment); (d) paperwork review (engines, APU); (e) Contractor travel and accommodation included in the Price.')
    add_para(doc, '1.2. Location: JAT Tehnika, Belgrade, Serbia. The Customer shall provide airport and Aircraft access.')
    add_para(doc, '1.3. Deliverables: inspection / BSI report(s) and BSI video materials. BSI shall be performed by a certified engineer.')

    add_para(doc, '2. TERM', bold=True, size=10.5, sb=2, sa=2, align='left')
    add_para(doc, '2.1. Approximate duration: 3 (three) days from the Effective Date and provision of access and required documentation by the Customer.')
    add_para(doc, '2.2. The Customer shall accept the report(s) within 5 (five) business days or submit reasoned objections; otherwise the Services are deemed accepted.')

    add_para(doc, '3. PRICE AND PAYMENT', bold=True, size=10.5, sb=2, sa=2, align='left')
    add_para(doc, '3.1. Total Price: USD 24,000.00 (twenty-four thousand US dollars), exclusive of taxes and fees:')
    for d, a in ITEMS_EN:
        add_para(doc, f'• {d} — {a}', align='left', sa=1)
    add_para(doc, 'Total: USD 24,000.00.', bold=True, sa=3, align='left')
    add_para(doc, '3.2. Payment option (select one; strike out the others):', sa=2)
    add_para(doc, '(A) 100% post-payment within 5 banking days after acceptance;', yellow_all=True, sa=1)
    add_para(doc, '(B) 50% on invoice and 50% within 5 banking days after acceptance; or', yellow_all=True, sa=1)
    add_para(doc, '(C) 100% advance within 3 banking days of invoice.', yellow_all=True, sa=1)
    add_para(doc, 'If not selected otherwise in writing, option (A) applies.', yellow_all=True, sa=2)
    add_para(doc, '3.3. Payment in USD to the Contractor\'s account in Clause 7. Remitting bank charges are for the Customer.')
    add_para(doc, '3.4. Extra work requires prior written agreement and separate pricing.')

    add_para(doc, '4. OBLIGATIONS', bold=True, size=10.5, sb=2, sa=2, align='left')
    add_para(doc, '4.1. The Contractor shall perform the Services professionally, comply with site safety rules and JAT Tehnika procedures, and deliver the reports and BSI video materials.')
    add_para(doc, '4.2. The Customer shall provide timely access, coordination with JAT Tehnika, required authorizations and documentation, payment, and a contact person.')
    add_para(doc, '4.3. The Services are inspection / diagnostic only and exclude repair, rectification, airworthiness certification and release to service unless agreed in writing.')

    add_para(doc, '5. LIABILITY AND FORCE MAJEURE', bold=True, size=10.5, sb=2, sa=2, align='left')
    add_para(doc, '5.1. Except for willful misconduct or gross negligence, the Contractor\'s liability is limited to the Price actually paid. Neither Party is liable for indirect or consequential losses.')
    add_para(doc, '5.2. Force majeure excuses performance while it continues, subject to prompt notice. If it exceeds 30 days, either Party may terminate without penalty, paying for Services already performed.')

    add_para(doc, '6. GENERAL', bold=True, size=10.5, sb=2, sa=2, align='left')
    add_para(doc, '6.1. This Agreement enters into force on signing (the "Effective Date") and remains valid until fully performed.')
    add_para(doc, '6.2. Either Party may terminate for material breach not cured within 10 business days. The Customer may cancel before commencement subject to reimbursement of documented costs.')
    add_para(doc, '6.3. Amendments require written form signed by both Parties (including PDF scans of signed originals).')
    add_para(doc, '6.4. Confidential information shall be protected for 3 years after termination.')
    add_para(doc, '6.5. Notices by email to Clause 7 addresses are deemed received on the next business day.')
    add_para(doc, '6.6. Governed by the laws of Türkiye. Disputes not settled within 15 business days shall be submitted to the courts of Istanbul, Türkiye (English unless otherwise agreed).')

    add_para(doc, '7. DETAILS AND SIGNATURES', bold=True, size=10.5, sb=2, sa=3, align='left')
    t = doc.add_table(rows=2, cols=2)
    fill_cell(t.rows[0].cells[0], f'CONTRACTOR: {TRADE}', [
        LEGAL, ADDR, f'Company No.: {CO_NO}   Tax No.: {TAX}',
        f'Bank: {BANK}', f'IBAN (USD): {IBAN_USD}', f'SWIFT: {SWIFT}',
        f'Tel.: {TEL}   Email: {EMAIL}',
    ])
    fill_cell(t.rows[0].cells[1], 'CUSTOMER: [●]', [
        'Name: [●]', 'Address: [●]', 'Reg. / Tax ID: [●]',
        'Bank / IBAN / SWIFT: [●]', 'Email: [●]   Tel.: [●]',
    ])
    fill_cell(t.rows[1].cells[0], '', ['', f'________________ / {DG} /', 'Director General', 'L.S.'])
    fill_cell(t.rows[1].cells[1], '', ['', '________________ / [Name] /', 'Title: [●]', 'L.S.'])
    add_para(doc, 'DRAFT — not legal advice.', size=8, sa=0, align='left')
    return doc


def build_ru():
    doc = Document()
    setup_doc(doc)
    add_para(doc, 'ДОГОВОР ВОЗМЕЗДНОГО ОКАЗАНИЯ УСЛУГ № ____ / JT-BSI-2026', bold=True, align='center', size=12, sa=1)
    add_para(doc, 'инспекция ВС и бороскопическая инспекция двигателей и ВСУ', bold=True, align='center', size=10.5, sa=1)
    add_para(doc, 'Boeing 737-500, регистрационный знак VQ-BOO', bold=True, align='center', size=10, sa=3)
    add_para(doc, 'Место заключения: ________________    Дата: «____» ____________ 2026 г.', align='center', size=9.5, sa=5)
    add_mixed(doc, [
        (f'{TRADE} / {LEGAL}', True),
        (f' ({ADDR}; Company No. {CO_NO}; Tax No. {TAX}; тел. {TEL}; {EMAIL}; «', False),
        ('Исполнитель', True),
        (f'»), в лице Генерального директора {DG}, с одной стороны, и', False),
    ], sa=3)
    add_mixed(doc, [
        ('[ПОЛНОЕ НАИМЕНОВАНИЕ ЗАКАЗЧИКА]', True),
        (' ([●]; рег. / налог. номер: [●]; «', False),
        ('Заказчик', True),
        ('»), в лице [ФИО, должность] на основании [Устава / доверенности], с другой стороны,', False),
    ], sa=3)
    add_para(doc, 'договорились о нижеследующем:', sa=4)

    add_para(doc, '1. ПРЕДМЕТ', bold=True, size=10.5, sa=2, align='left')
    add_para(doc, '1.1. Исполнитель оказывает, а Заказчик принимает и оплачивает следующие услуги в отношении Boeing 737-500 VQ-BOO (далее — «ВС»):')
    add_para(doc, '(a) BSI двух (2) двигателей CFM56-3 (объём уточняется); (b) BSI ВСУ APS2000; (c) GVI ВС (шасси, кабина, отсек авионики); (d) проверка документации (двигатели, ВСУ); (e) транспорт и проживание Исполнителя в объёме, включённом в Цену.')
    add_para(doc, '1.2. Место оказания: JAT Tehnika, Белград, Сербия. Заказчик обеспечивает доступ в аэропорт и к ВС.')
    add_para(doc, '1.3. Результат: отчёт(ы) об инспекции / BSI и видеоматериалы BSI. BSI выполняет сертифицированный инженер.')

    add_para(doc, '2. СРОК', bold=True, size=10.5, sb=2, sa=2, align='left')
    add_para(doc, '2.1. Ориентировочно 3 (три) дня с даты вступления Договора в силу и предоставления Заказчиком доступа и необходимой документации.')
    add_para(doc, '2.2. Заказчик принимает отчёт(ы) в течение 5 (пяти) рабочих дней либо направляет мотивированные возражения; иначе Услуги считаются принятыми.')

    add_para(doc, '3. ЦЕНА И ОПЛАТА', bold=True, size=10.5, sb=2, sa=2, align='left')
    add_para(doc, '3.1. Общая Цена: USD 24 000,00 (двадцать четыре тысячи долларов США) без учёта налогов и сборов:')
    for d, a in ITEMS_RU:
        add_para(doc, f'• {d} — {a}', align='left', sa=1)
    add_para(doc, 'Итого: USD 24 000,00.', bold=True, sa=3, align='left')
    add_para(doc, '3.2. Вариант оплаты (выбрать один; остальные вычеркнуть):', sa=2)
    add_para(doc, '(A) 100% постоплата в течение 5 банковских дней после приёмки;', yellow_all=True, sa=1)
    add_para(doc, '(B) 50% по счёту и 50% в течение 5 банковских дней после приёмки; либо', yellow_all=True, sa=1)
    add_para(doc, '(C) 100% предоплата в течение 3 банковских дней с даты счёта.', yellow_all=True, sa=1)
    add_para(doc, 'Если иное не выбрано письменно, применяется вариант (A).', yellow_all=True, sa=2)
    add_para(doc, '3.3. Оплата в USD на счёт Исполнителя (раздел 7). Комиссии банка-отправителя — за счёт Заказчика.')
    add_para(doc, '3.4. Дополнительные работы — только по предварительному письменному согласованию и отдельной цене.')

    add_para(doc, '4. ОБЯЗАННОСТИ', bold=True, size=10.5, sb=2, sa=2, align='left')
    add_para(doc, '4.1. Исполнитель оказывает Услуги профессионально, соблюдает правила ТБ и процедуры JAT Tehnika, передаёт отчёт(ы) и видеоматериалы BSI.')
    add_para(doc, '4.2. Заказчик своевременно обеспечивает доступ, координацию с JAT Tehnika, разрешения и документацию, оплату и контактное лицо.')
    add_para(doc, '4.3. Услуги носят инспекционный характер и не включают ремонт, устранение дефектов, подтверждение лётной годности и допуск к эксплуатации, если иное не согласовано письменно.')

    add_para(doc, '5. ОТВЕТСТВЕННОСТЬ И ФОРС-МАЖОР', bold=True, size=10.5, sb=2, sa=2, align='left')
    add_para(doc, '5.1. За исключением умысла или грубой неосторожности ответственность Исполнителя ограничена фактически уплаченной Ценой. Стороны не отвечают за косвенные убытки.')
    add_para(doc, '5.2. Форс-мажор освобождает от ответственности на период его действия при своевременном уведомлении. При продолжении более 30 дней любая Сторона вправе расторгнуть Договор без штрафов с оплатой фактически оказанных Услуг.')

    add_para(doc, '6. ОБЩИЕ ПОЛОЖЕНИЯ', bold=True, size=10.5, sb=2, sa=2, align='left')
    add_para(doc, '6.1. Договор вступает в силу с даты подписания и действует до полного исполнения.')
    add_para(doc, '6.2. Любая Сторона вправе расторгнуть Договор при существенном нарушении, не устранённом в течение 10 рабочих дней. Заказчик вправе отказаться до начала работ с возмещением документально подтверждённых расходов Исполнителя.')
    add_para(doc, '6.3. Изменения — только в письменной форме, подписанной обеими Сторонами (включая PDF-сканы подписанных оригиналов).')
    add_para(doc, '6.4. Конфиденциальная информация защищается в течение 3 лет после прекращения Договора.')
    add_para(doc, '6.5. Уведомления по email (раздел 7) считаются полученными на следующий рабочий день.')
    add_para(doc, '6.6. Применимое право — Турция. Споры, не урегулированные за 15 рабочих дней, передаются в суды Стамбула (язык — английский, если не согласовано иное).')

    add_para(doc, '7. РЕКВИЗИТЫ И ПОДПИСИ', bold=True, size=10.5, sb=2, sa=3, align='left')
    t = doc.add_table(rows=2, cols=2)
    fill_cell(t.rows[0].cells[0], f'ИСПОЛНИТЕЛЬ: {TRADE}', [
        LEGAL, ADDR, f'Company No.: {CO_NO}   Tax No.: {TAX}',
        f'Банк: {BANK}', f'IBAN (USD): {IBAN_USD}', f'SWIFT: {SWIFT}',
        f'Тел.: {TEL}   Email: {EMAIL}',
    ])
    fill_cell(t.rows[0].cells[1], 'ЗАКАЗЧИК: [●]', [
        'Наименование: [●]', 'Адрес: [●]', 'Рег. / налог. номер: [●]',
        'Банк / IBAN / SWIFT: [●]', 'Email: [●]   Tel.: [●]',
    ])
    fill_cell(t.rows[1].cells[0], '', ['', f'________________ / {DG} /', 'Генеральный директор', 'М.П.'])
    fill_cell(t.rows[1].cells[1], '', ['', '________________ / [ФИО] /', 'Должность: [●]', 'М.П.'])
    add_para(doc, 'ЧЕРНОВИК — не юридическая консультация.', size=8, sa=0, align='left')
    return doc


def patch_rfonts_element(rfonts, font=FONT):
    theme_keys = [k for k in list(rfonts.attrib.keys()) if 'Theme' in k]
    for k in theme_keys:
        del rfonts.attrib[k]
    for tag in ('ascii', 'hAnsi', 'cs', 'eastAsia'):
        rfonts.set(f'{W}{tag}', font)


def patch_xml_root(root, font=FONT):
    for rfonts in root.iter(f'{W}rFonts'):
        patch_rfonts_element(rfonts, font)
    for r in root.iter(f'{W}r'):
        texts = ''.join(t.text or '' for t in r.iter(f'{W}t'))
        if not texts:
            continue
        rpr = r.find(f'{W}rPr')
        if rpr is None:
            rpr = ET.Element(f'{W}rPr')
            r.insert(0, rpr)
        rf = rpr.find(f'{W}rFonts')
        if rf is None:
            rf = ET.SubElement(rpr, f'{W}rFonts')
        patch_rfonts_element(rf, font)
    for latin in root.iter(f'{A}latin'):
        latin.set('typeface', font)


def patch_docx_fonts(path: Path, font: str = FONT):
    replacements = (
        ('typeface="Cambria"', f'typeface="{font}"'),
        ('typeface="Times New Roman"', f'typeface="{font}"'),
        ('w:ascii="Times New Roman"', f'w:ascii="{font}"'),
        ('w:hAnsi="Times New Roman"', f'w:hAnsi="{font}"'),
        ('w:cs="Times New Roman"', f'w:cs="{font}"'),
        ('w:ascii="Cambria"', f'w:ascii="{font}"'),
        ('w:hAnsi="Cambria"', f'w:hAnsi="{font}"'),
        ('w:cs="Cambria"', f'w:cs="{font}"'),
    )
    buf = io.BytesIO()
    with zipfile.ZipFile(path, 'r') as zin, zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith('word/') and item.filename.endswith('.xml'):
                text = data.decode('utf-8')
                for old, new in replacements:
                    text = text.replace(old, new)
                root = ET.fromstring(text.encode('utf-8'))
                patch_xml_root(root, font)
                data = ET.tostring(root, encoding='utf-8', xml_declaration=True)
            zout.writestr(item, data)
    path.write_bytes(buf.getvalue())


def audit_docx(path: Path):
    with zipfile.ZipFile(path) as z:
        fonts = {}
        for name in z.namelist():
            if not (name.startswith('word/') and name.endswith('.xml')):
                continue
            root = ET.fromstring(z.read(name))
            for rf in root.iter(f'{W}rFonts'):
                for k, v in rf.attrib.items():
                    if 'Theme' in k:
                        fonts[f'THEME:{v.split("}")[-1]}'] = fonts.get(f'THEME:{v.split("}")[-1]}', 0) + 1
                    elif 'ascii' in k or 'hAnsi' in k:
                        fonts[v] = fonts.get(v, 0) + 1
        print(path.name, fonts)


def main():
    outputs = [
        (build_en(), DIR / 'Service_Agreement_JT-BSI-2026_EN.docx'),
        (build_ru(), DIR / 'Dogovor_JT-BSI-2026_RU.docx'),
    ]
    for doc, path in outputs:
        doc.save(path)
        patch_docx_fonts(path)
        audit_docx(path)


if __name__ == '__main__':
    main()
